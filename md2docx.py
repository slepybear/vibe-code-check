#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, RGBColor
import re
import sys

def md_to_docx(md_path, docx_path):
    doc = Document()
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    for line in lines:
        line = line.rstrip()
        
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
        elif line.startswith('```'):
            continue
        elif line.strip() == '':
            doc.add_paragraph('')
        elif line.startswith('---'):
            p = doc.add_paragraph('_' * 50)
            p.paragraph_format.alignment = 1
        else:
            # 处理粗体
            original_line = line
            line = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
            # 处理链接
            line = re.sub(r'\[(.*?)\]\((.*?)\)', r'\1: \2', line)
            
            # 处理列表
            if line.startswith('- ') or line.startswith('* '):
                p = doc.add_paragraph(line[2:], style='List Bullet')
            elif line.strip() and line[0].isdigit() and line[1] in ('.', ')'):
                p = doc.add_paragraph(line, style='List Number')
            else:
                p = doc.add_paragraph(line)
    
    doc.save(docx_path)
    print(f'成功转换: {docx_path}')

if __name__ == '__main__':
    try:
        md_to_docx('AI_AGENT_JOURNEY.md', 'AI_AGENT_JOURNEY.docx')
    except ImportError:
        print('需要安装 python-docx: pip install python-docx')
        import subprocess
        subprocess.run([sys.executable, '-m', 'pip', 'install', 'python-docx'])
        md_to_docx('AI_AGENT_JOURNEY.md', 'AI_AGENT_JOURNEY.docx')

